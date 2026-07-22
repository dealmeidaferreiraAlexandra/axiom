import pytest
from docx import Document
from pypdf import PdfWriter

from core.loader import (
    load_docx,
    load_file,
    load_file_pages,
    load_pdf,
    load_txt,
)

reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")


def _make_text_pdf(path, pages):
    """Create a PDF whose pages contain the given lines of extractable text."""
    c = reportlab_canvas.Canvas(str(path))
    for lines in pages:
        for i, line in enumerate(lines):
            c.drawString(72, 720 - i * 20, line)
        c.showPage()
    c.save()


def test_load_txt_reads_content(tmp_path):
    f = tmp_path / "sample.txt"
    f.write_text("hello from a text file", encoding="utf-8")
    assert load_txt(str(f)) == "hello from a text file"


def test_load_txt_handles_unicode(tmp_path):
    f = tmp_path / "unicode.txt"
    f.write_text("olá acentuação çãõ", encoding="utf-8")
    assert load_txt(str(f)) == "olá acentuação çãõ"


def test_load_docx_joins_paragraphs(tmp_path):
    f = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("first paragraph")
    doc.add_paragraph("second paragraph")
    doc.save(str(f))
    assert load_docx(str(f)) == "first paragraph\nsecond paragraph"


def test_load_pdf_missing_file_returns_empty_string():
    assert load_pdf("/does/not/exist.pdf") == ""


def test_load_pdf_blank_pages_return_empty_string(tmp_path):
    f = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(f, "wb") as out:
        writer.write(out)
    # A blank page has no extractable text.
    assert load_pdf(str(f)) == ""


def test_load_pdf_extracts_text_from_pages(tmp_path):
    f = tmp_path / "text.pdf"
    _make_text_pdf(f, [["First page line"], ["Second page line"]])
    text = load_pdf(str(f))
    assert "First page line" in text
    assert "Second page line" in text


def test_load_file_dispatches_pdf(tmp_path):
    f = tmp_path / "doc.pdf"
    _make_text_pdf(f, [["Dispatched via load_file"]])
    assert "Dispatched via load_file" in load_file(str(f))


def test_load_file_dispatches_docx(tmp_path):
    f = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("docx via load_file")
    doc.save(str(f))
    assert load_file(str(f)) == "docx via load_file"


def test_load_file_pages_pdf_returns_numbered_pages(tmp_path):
    f = tmp_path / "multi.pdf"
    _make_text_pdf(f, [["Alpha page"], ["Bravo page"]])
    pages = load_file_pages(str(f))
    numbers = [n for n, _ in pages]
    assert numbers == [1, 2]
    assert "Alpha page" in pages[0][1]
    assert "Bravo page" in pages[1][1]


def test_load_file_dispatches_by_extension(tmp_path):
    txt = tmp_path / "note.TXT"  # also verifies case-insensitivity
    txt.write_text("dispatch me", encoding="utf-8")
    assert load_file(str(txt)) == "dispatch me"


def test_load_file_unknown_extension_returns_empty_string(tmp_path):
    other = tmp_path / "data.csv"
    other.write_text("a,b,c", encoding="utf-8")
    assert load_file(str(other)) == ""


def test_load_file_pages_for_txt_returns_single_none_page(tmp_path):
    f = tmp_path / "note.txt"
    f.write_text("page-less content", encoding="utf-8")
    assert load_file_pages(str(f)) == [(None, "page-less content")]


def test_load_file_pages_empty_txt_returns_no_pages(tmp_path):
    f = tmp_path / "empty.txt"
    f.write_text("   \n  ", encoding="utf-8")
    assert load_file_pages(str(f)) == []


def test_load_file_pages_missing_pdf_returns_empty_list():
    assert load_file_pages("/does/not/exist.pdf") == []
